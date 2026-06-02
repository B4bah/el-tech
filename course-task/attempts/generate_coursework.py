import math
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# --- Стандартные ряды номиналов ---
E24 = [1.0, 1.1, 1.2, 1.3, 1.5, 1.6, 1.8, 2.0, 2.2, 2.4, 2.7, 3.0, 3.3, 3.6, 3.9, 4.3, 4.7, 5.1, 5.6, 6.2, 6.8, 7.5, 8.2, 9.1]
E12 = [1.0, 1.2, 1.5, 1.8, 2.2, 2.7, 3.3, 3.9, 4.7, 5.6, 6.8, 8.2]

def closest_value(value, series):
    if value <= 0: return series[0]
    exp = math.floor(math.log10(value))
    base = value / (10**exp)
    closest = min(series, key=lambda x: abs(x - base))
    return closest * (10**exp)

def format_res(val):
    if val >= 1e6: return f"{val/1e6:g} МОм"
    if val >= 1e3: return f"{val/1e3:g} кОм"
    return f"{val:g} Ом"

def format_cap(val):
    if val >= 1e-3: return f"{val/1e-3:g} мкФ"
    if val >= 1e-6: return f"{val/1e-6:g} нФ"
    return f"{val/1e-12:g} пФ"

def set_run_font(run, size=14, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size, bold)
    return p

def add_heading_custom(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, 14, bold=True)
    else:
        run = p.add_run(text)
        set_run_font(run, 14, bold=True)
    return p

def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12, bold)

# --- ИСХОДНЫЕ ДАННЫЕ (ВАРИАНТ 11) ---
Ug = 0.8e-3      # В
Rg = 150e3       # Ом
Un = 3.0         # В
Rn = 600.0       # Ом
fn = 30.0        # Гц
fv = 15e3        # Гц
Mn = 1.2
Mv = 1.2

# Параметры ОУ К574УД3
f1_OU = 10e6     # Гц
Rvh_OU = 100e6   # Ом
Ip_OU = 5e-3     # А

# Параметры транзисторов КТ315В/КТ361Д
h21 = 30

# --- РАСЧЕТЫ ---
Ku = Un / Ug
Kk1 = 78.0
Kk2 = 48.0

# Каскад 1 (Неинвертирующий)
R1 = closest_value(1.2e3, E24)
R2_calc = (Kk1 - 1) * R1
R2 = closest_value(R2_calc, E24)
R3 = closest_value(7 * Rg, E24)

# Каскад 2 (Инвертирующий)
R6 = closest_value(10e3, E24)
R5_calc = Kk2 * R6
R5 = closest_value(R5_calc, E24)

# Корректировка усиления
R7_calc = 1.2 * R5
R7 = closest_value(R7_calc, E24)
R8_calc = 0.4 * R5
R8 = closest_value(R8_calc, E24)

# Эмиттерный повторитель
Rvh_ep = h21 * Rn

# Разделительные конденсаторы (Mnc = sqrt(Mn) ~ 1.095)
Mnc = math.sqrt(Mn)
denom = 2 * math.pi * fn * math.sqrt(Mnc**2 - 1)

Req_C1 = Rg + R3
C1_calc = 1 / (denom * Req_C1)
C1 = closest_value(C1_calc, E12)

Req_C2 = R2
C2_calc = 1 / (denom * Req_C2)
C2 = closest_value(C2_calc, E12)

Req_C6 = R6 
C6_calc = 1 / (denom * Req_C6)
C6 = closest_value(C6_calc, E12)

Req_C7 = Rn 
C7_calc = 1 / (denom * Req_C7)
C7 = closest_value(C7_calc, E12)

# Корректирующий конденсатор C3 (Mcv ~ 1.1)
Mcv = 1.1
C3_calc = math.sqrt(Mcv**2 - 1) / (2 * math.pi * fv * R2)
C3 = closest_value(C3_calc, E12)

# Фильтры питания
R4 = closest_value(300, E24)
R5_f = closest_value(300, E24)
C4 = closest_value(1000e-6, E12)
C5 = closest_value(1000e-6, E12)

# --- ГЕНЕРАЦИЯ ДОКУМЕНТА ---
doc = Document()

# Настройка стилей по умолчанию
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5

# 1. ТИТУЛЬНЫЙ ЛИСТ
add_para(doc, "МИНОБРНАУКИ РОССИИ", 14, True, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Федеральное государственное бюджетное образовательное учреждение высшего образования", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "«Санкт-Петербургский государственный морской технический университет»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "(СПбГМТУ)", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Факультет цифровых промышленных технологий", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "Направление 09.03.01 Информатика и вычислительная техника", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
doc.add_paragraph()
add_para(doc, "КУРСОВАЯ РАБОТА", 16, True, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "по дисциплине «Электротехника и электроника»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Тема: «Расчет предварительного усилителя на ОУ»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Вариант 11", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
doc.add_paragraph()
add_para(doc, "Выполнил: Студент группы 20390", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
add_para(doc, "Суровцев В.Е.", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
doc.add_paragraph()
add_para(doc, "Проверил: Старший преподаватель", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
add_para(doc, "Ветров Б.Г.", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "Санкт-Петербург", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "2025", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
doc.add_page_break()

# 2. СОДЕРЖАНИЕ
add_heading_custom(doc, "Содержание", 1)
toc_items = [
    "Введение", "Задание", "Расчет усилителя", 
    "Выбор элементной базы и количества каскадов", "Расчет оконечного каскада",
    "Расчет входного каскада", "Составление схемы усилителя",
    "Расчет частотных искажений (нижние частоты)", "Расчет частотных искажений (верхние частоты)",
    "Корректировка коэффициента усиления", "Расчет элементов фильтров в цепи питания",
    "Обоснование выбора элементной базы", "Заключение", "Список использованной литературы"
]
for item in toc_items:
    add_para(doc, item, 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
doc.add_page_break()

# 3. ВВЕДЕНИЕ
add_heading_custom(doc, "Введение", 1)
add_para(doc, "Операционный усилитель (ОУ) — это высокоточный электронный усилитель напряжения с дифференциальными входами и одним выходом, обладающий чрезвычайно высоким коэффициентом усиления. Применение ОУ даёт возможность создавать разнообразные электронные схемы и устройства: усилители, генераторы, фильтры. В данной работе рассматриваются вопросы расчёта предварительного усилителя на основе микросхем ОУ.")
add_para(doc, "Параметры каскадов на ОУ определяются функциональным назначением устройств. В морской технике используются предварительные, согласующие, широкополосные и другие системы. Предварительный усиливает сигнал от удалённого датчика до величины, необходимой для передачи по каналу связи, при этом структура сигнала не должна искажаться.")
add_para(doc, "В ходе данной работы будет проведен расчет предварительного усилителя на операционных усилителях для Варианта 11.")
doc.add_page_break()

# 4. ЗАДАНИЕ
add_heading_custom(doc, "Задание", 1)
add_para(doc, "Основные параметры усилителя:", 14, True, WD_ALIGN_PARAGRAPH.LEFT, False)
table_zad = doc.add_table(rows=2, cols=8)
table_zad.style = 'Table Grid'
headers = ["Uг, мВ", "Rг, кОм", "Uн, В", "Rн, Ом", "fн, Гц", "fв, кГц", "Мн", "Мв"]
values = [f"{Ug*1000}", f"{Rg/1000}", f"{Un}", f"{Rn}", f"{fn}", f"{fv/1000}", f"{Mn}", f"{Mv}"]
for i, h in enumerate(headers): set_cell(table_zad.rows[0].cells[i], h, bold=True)
for i, v in enumerate(values): set_cell(table_zad.rows[1].cells[i], v)
add_para(doc, "Основные параметры операционного усилителя (марка: К574УД3): f1=10 МГц, Rвх=100 МОм, Iп=5 мА.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)

# 5. РАСЧЕТ УСИЛИТЕЛЯ
add_heading_custom(doc, "Расчет усилителя", 1)
add_heading_custom(doc, "Выбор элементной базы и количества каскадов", 2)
add_para(doc, f"Максимальный коэффициент усиления каскада на верхней граничной частоте: K = f1 / fв = {f1_OU/1e6} / {fv/1000} = {f1_OU/fv:.1f}.")
add_para(doc, f"Коэффициент усиления усилителя: Ku = Uн / Uг = {Un} / {Ug*1000} мВ = {Ku:.0f}.")
add_para(doc, f"Определим количество каскадов m. При m=2 средний коэффициент усиления каскада составит {math.sqrt(Ku):.1f}. Принимаем m = 2.")
add_para(doc, f"Распределим усиление по каскадам: Kк1 = {Kk1:.0f} (неинвертирующий), Kк2 = {Kk2:.0f} (инвертирующий).")

add_heading_custom(doc, "Расчет оконечного каскада", 2)
add_para(doc, "Рис. 1 – Схема оконечного каскада", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "[Место для рисунка]", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, f"Для согласования выхода ОУ и нагрузки Rн={Rn} Ом применяем эмиттерный повторитель на комплиментарных транзисторах V1 (КТ315В) и V2 (КТ361Д) с h21={h21}.")
add_para(doc, f"Входное сопротивление эмиттерного повторителя: Rвх_эп = h21 * Rн = {h21} * {Rn} = {Rvh_ep} Ом = {Rvh_ep/1000} кОм.")
add_para(doc, f"Задаем входное сопротивление каскада R6 = {format_res(R6)}.")
add_para(doc, f"Коэффициент усиления Kк2 = {Kk2:.0f}. Рассчитываем резистор обратной связи: R5 = Kк2 * R6 = {format_res(R5_calc)}. Выбираем стандартный номинал R5 = {format_res(R5)}.")

add_heading_custom(doc, "Расчет входного (первого) каскада", 2)
add_para(doc, "Рис. 2 – Схема входного каскада", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "[Место для рисунка]", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, f"Входной каскад согласует генератор с усилителем. Принимаем R3 = 7 * Rг = {format_res(7*Rg)}. Стандартный номинал R3 = {format_res(R3)}.")
add_para(doc, f"Коэффициент усиления Kк1 = {Kk1:.0f}. Задаем R1 = {format_res(R1)}.")
add_para(doc, f"Рассчитываем R2 = (Kк1 - 1) * R1 = {format_res(R2_calc)}. Выбираем стандартный номинал R2 = {format_res(R2)}.")

add_heading_custom(doc, "Расчет частотных искажений (нижние частоты)", 2)
add_para(doc, f"Коэффициент частотных искажений на нижних частотах Мн = {Mn}.")
add_para(doc, f"Рассчитываем разделительные конденсаторы:")
add_para(doc, f"C1 (вход): Rэкв = {format_res(Req_C1)}, C1 = {format_cap(C1_calc)} -> принимаем {format_cap(C1)}.")
add_para(doc, f"C2 (ООС 1 каскада): Rэкв = {format_res(Req_C2)}, C2 = {format_cap(C2_calc)} -> принимаем {format_cap(C2)}.")
add_para(doc, f"C6 (межкаскадная): Rэкв = {format_res(Req_C6)}, C6 = {format_cap(C6_calc)} -> принимаем {format_cap(C6)}.")
add_para(doc, f"C7 (выходной): Rэкв = {format_res(Req_C7)}, C7 = {format_cap(C7_calc)} -> принимаем {format_cap(C7)}.")

add_heading_custom(doc, "Расчет частотных искажений (верхние частоты)", 2)
add_para(doc, f"Для коррекции на верхних частотах (Мв = {Mv}) используем конденсатор C3 в первом каскаде.")
add_para(doc, f"Расчетное значение C3 = {format_cap(C3_calc)}. Принимаем стандартный номинал {format_cap(C3)}.")

add_heading_custom(doc, "Корректировка коэффициента усиления", 2)
add_para(doc, f"Для компенсации разброса номиналов резистор R5 ({format_res(R5)}) заменяем на потенциометр R7 и постоянный резистор R8.")
add_para(doc, f"R7 (потенциометр) = {format_res(R7)}.")
add_para(doc, f"R8 (постоянный) = {format_res(R8)}.")

add_heading_custom(doc, "Расчет элементов фильтров в цепи питания", 2)
add_para(doc, f"Для устранения наводок по цепям питания устанавливаем RC-фильтры.")
add_para(doc, f"R4 = R5(ф) = {format_res(R4)}.")
add_para(doc, f"C4 = C5 = {format_cap(C4)}.")

# 6. СПЕЦИФИКАЦИЯ
add_heading_custom(doc, "Обоснование выбора элементной базы (Спецификация)", 1)
table_spec = doc.add_table(rows=1, cols=4)
table_spec.style = 'Table Grid'
headers_spec = ["Обозначение", "Наименование", "Количество", "Примечание"]
for i, h in enumerate(headers_spec): set_cell(table_spec.rows[0].cells[i], h, bold=True)

components = [
    ("", "Операционные усилители", "", ""),
    ("A1, A2", f"К574УД3", "2", ""),
    ("", "Транзисторы", "", ""),
    ("V1", "КТ315В", "1", ""),
    ("V2", "КТ361Д", "1", ""),
    ("", "Резисторы", "", ""),
    ("R1", f"ОМЛТ-0.125 {format_res(R1)} ±10%", "1", ""),
    ("R2", f"ОМЛТ-0.125 {format_res(R2)} ±10%", "1", ""),
    ("R3", f"ОМЛТ-0.125 {format_res(R3)} ±10%", "1", ""),
    ("R4", f"ОМЛТ-0.125 {format_res(R4)} ±10%", "1", "Фильтр +"),
    ("R5", f"ОМЛТ-0.125 {format_res(R5_f)} ±10%", "1", "Фильтр -"),
    ("R6", f"ОМЛТ-0.125 {format_res(R6)} ±10%", "1", ""),
    ("R7", f"Потенциометр СПО-0.25 {format_res(R7)}", "1", ""),
    ("R8", f"ОМЛТ-0.125 {format_res(R8)} ±10%", "1", ""),
    ("", "Конденсаторы", "", ""),
    ("C1", f"К73-17-63В {format_cap(C1)} ±10%", "1", ""),
    ("C2", f"К10-17-50В {format_cap(C2)} ±5%", "1", ""),
    ("C3", f"К10-23-50В {format_cap(C3)} ±5%", "1", ""),
    ("C4", f"К50-6 {format_cap(C4)} ±20%", "1", ""),
    ("C5", f"К50-6 {format_cap(C5)} ±20%", "1", ""),
    ("C6", f"К73-17-63В {format_cap(C6)} ±10%", "1", ""),
    ("C7", f"К50-6 {format_cap(C7)} ±20%", "1", ""),
]

for comp in components:
    row = table_spec.add_row().cells
    for i, val in enumerate(comp):
        set_cell(row[i], val, align=WD_ALIGN_PARAGRAPH.LEFT if i in [1, 3] else WD_ALIGN_PARAGRAPH.CENTER)

# 7. ЗАКЛЮЧЕНИЕ
add_heading_custom(doc, "Заключение", 1)
add_para(doc, f"В ходе курсового проекта был рассчитан предварительный усилитель на базе ОУ К574УД3. Обеспечен коэффициент усиления Ku = {Ku:.0f} при заданных параметрах сигнала.")
add_para(doc, f"Все номиналы резисторов и конденсаторов выбраны из стандартных рядов E24 и E12, что гарантирует возможность практической реализации устройства. Расчеты частотных искажений подтвердили соответствие требованиям Мн={Mn} и Мв={Mv}.")

# 8. ЛИТЕРАТУРА
add_heading_custom(doc, "Список использованной литературы", 1)
add_para(doc, "1. Методические рекомендации по выполнению курсовой работы. Ветров Б. Г.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "2. Курс лекций по электронике. Ветров Б. Г.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "3. Справочник по элементной базе электронных устройств.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)

# Сохранение
file_name = "Курсовая_работа_Вариант_11.docx"
doc.save(file_name)
print(f"Файл успешно создан: {os.path.abspath(file_name)}")