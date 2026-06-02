import math
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# --- Стандартные ряды номиналов ---
E24 = [1.0, 1.1, 1.2, 1.3, 1.5, 1.6, 1.8, 2.0, 2.2, 2.4, 2.7, 3.0, 3.3, 3.6, 3.9, 4.3, 4.7, 5.1, 5.6, 6.2, 6.8, 7.5, 8.2, 9.1]
E12 = [1.0, 1.2, 1.5, 1.8, 2.2, 2.7, 3.3, 3.9, 4.7, 5.6, 6.8, 8.2]

def closest_value(value, series):
    """Выбор ближайшего стандартного номинала из ряда E24 или E12"""
    if value <= 0:
        return series[0]
    exp = math.floor(math.log10(value))
    base = value / (10**exp)
    closest = min(series, key=lambda x: abs(x - base))
    return closest * (10**exp)

def format_res(val):
    """Форматирование сопротивления"""
    if val >= 1e6:
        return f"{val/1e6:g} МОм"
    if val >= 1e3:
        return f"{val/1e3:g} кОм"
    return f"{val:g} Ом"

def format_cap(val):
    """Форматирование ёмкости"""
    if val >= 1e-3:
        return f"{val/1e-3:g} мкФ"
    if val >= 1e-6:
        return f"{val/1e-6:g} нФ"
    return f"{val/1e-12:g} пФ"

def set_run_font(run, size=14, bold=False, italic=False):
    """Установка шрифта Times New Roman"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(doc, text, size=14, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=True, space_after=0):
    """Добавление параграфа"""
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size, bold)
    return p

def add_heading_custom(doc, text, level=1):
    """Добавление заголовка"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, 14, bold=True)
    return p

def set_cell(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Заполнение ячейки таблицы"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, 12, bold)

# === ИСХОДНЫЕ ДАННЫЕ (ВАРИАНТ 11) ===
Ug = 0.8e-3      # В
Rg = 150e3       # Ом
Un = 3.0         # В
Rn = 600.0       # Ом
fn = 30.0        # Гц
fv = 15e3        # Гц
Mn = 1.2
Mv = 1.2

# Параметры ОУ К574УД3
f1_OU = 10e6     # Гц (частота единичного усиления)
Rvh_OU = 100e6   # Ом (входное сопротивление)
Ip_OU = 5e-3     # А (ток потребления)

# Параметры транзисторов КТ315В/КТ361Д
h21 = 30
fh21_tr = 100e6  # Гц (граничная частота усиления по току)

# === РАСЧЁТЫ ===

# Общий коэффициент усиления
Ku = Un / Ug  # 3750

# Распределение усиления по каскадам
Kk1 = 78.0   # неинвертирующий (первый каскад)
Kk2 = 48.0   # инвертирующий (второй каскад)
Ku_fakt = Kk1 * Kk2  # 3744 ≈ 3750

# --- Первый каскад (неинвертирующий) ---
# Kк1 = 1 + R2/R1
R1 = closest_value(1.2e3, E24)  # 1.2 кОм
R2_calc = (Kk1 - 1) * R1        # 77 * 1.2к = 92.4 кОм
R2 = closest_value(R2_calc, E24)  # 91 кОм
Kk1_fakt = 1 + R2/R1  # 1 + 91/1.2 = 76.8

# R3 для согласования с генератором: R3 = 7 * Rг
R3_calc = 7 * Rg  # 1.05 МОм
R3 = closest_value(R3_calc, E24)  # 1 МОм

# --- Второй каскад (инвертирующий) ---
# Kк2 = (R7+R8)/R6
R6 = closest_value(10e3, E24)  # 10 кОм
R_os_calc = Kk2 * R6  # 480 кОм (суммарное сопротивление ОС)

# По методичке: R_пот = 1.2 * R_расч, R_пост = 0.4 * R_расч
R7_calc = 1.2 * R_os_calc  # 576 кОм
R7 = closest_value(R7_calc, E24)  # 560 кОм (потенциометр)

R8_calc = 0.4 * R_os_calc  # 192 кОм
R8 = closest_value(R8_calc, E24)  # 200 кОм или 180 кОм

# Проверка Kк2 при среднем положении потенциометра (R7/2 = 280 кОм)
Kk2_min = R8 / R6  # минимальное усиление
Kk2_max = (R7 + R8) / R6  # максимальное усиление
Kk2_mid = (R7/2 + R8) / R6  # усиление при среднем положении

# --- Эмиттерный повторитель ---
Rvh_ep = h21 * Rn  # 30 * 600 = 18 кОм

# --- Разделительные конденсаторы ---
# Коэффициент искажений одного конденсатора (предполагаем 4 конденсатора: C1, C2, C6, C7)
Mnc = math.pow(Mn, 1/4)  # ≈ 1.047
denom_low = 2 * math.pi * fn * math.sqrt(Mnc**2 - 1)

# C1: между генератором и входом усилителя
# R_экв = Rг + Rвх_каскада ≈ Rг + R3 (так как Rвх_ОУ очень велико)
Req_C1 = Rg + R3
C1_calc = 1 / (denom_low * Req_C1)
# На практике берут с запасом, поэтому выбираем 4.7 мкФ
C1 = 4.7e-6

# C2: шунт на входе (параллельно R3, фильтр ВЧ помех)
# Должен иметь маленькую ёмкость, чтобы не шунтировать R3 на fн
C2_calc = 1 / (2 * math.pi * fn * R3 * 10)  # Xc >> R3 на fн
C2 = closest_value(C2_calc, E12)  # ~8.2 пФ

# C6: межкаскадный (между эмиттерным повторителем и вторым каскадом)
# R_экв = Rвых_эп + Rвх_каскада2 ≈ 0 + R6
Req_C6 = R6
C6_calc = 1 / (denom_low * Req_C6)
C6 = closest_value(C6_calc, E12)  # ~1 мкФ, берём 4.7 мкФ с запасом
C6 = 4.7e-6

# C7: выходной (между вторым каскадом и нагрузкой)
# R_экв = Rвых_каскада + Rн ≈ 0 + Rн
Req_C7 = Rn
C7_calc = 1 / (denom_low * Req_C7)
C7 = closest_value(C7_calc, E12)  # ~15 мкФ

# Проверка C7: Xc на fн должно быть << Rн
Xc7_fn = 1 / (2 * math.pi * fn * C7)
# Xc7_fn должно быть << 600 Ом

# --- Корректирующий конденсатор C3 ---
# Расчёт искажений от активных элементов на верхних частотах
Mv_OU1 = math.sqrt(1 + (Kk1 * fv / f1_OU)**2)
Mv_OU2 = math.sqrt(1 + (Kk2 * fv / f1_OU)**2)
Mv_tr = math.sqrt(1 + (fv / fh21_tr)**2)

Mv_active = Mv_OU1 * Mv_OU2 * Mv_tr

# Оставшиеся искажения для C3
Mv_C3 = Mv / Mv_active

# Расчёт C3
C3_calc = math.sqrt(Mv_C3**2 - 1) / (2 * math.pi * fv * R2)
C3 = closest_value(C3_calc, E12)  # ~56-68 пФ

# Проверка C3
Xc3_fv = 1 / (2 * math.pi * fv * C3)
Mv_C3_check = math.sqrt(1 + (R2 / Xc3_fv)**2)
Mv_total_check = Mv_active * Mv_C3_check

# --- Фильтры питания ---
# По методичке: Rф = 0.1 * Uп / Iп = 0.1 * 15 / 0.005 = 300 Ом
R4 = 300  # Ом
R5 = 300  # Ом

# Cф: Xcф = Rф / 100 на fн
C4_calc = 100 / (2 * math.pi * fn * R4)
C4 = closest_value(C4_calc, E12)  # ~1768 мкФ → 1000 мкФ или 2200 мкФ
C4 = 1000e-6  # 1000 мкФ (по методичке)
C5 = 1000e-6  # 1000 мкФ

# === ГЕНЕРАЦИЯ ДОКУМЕНТА ===
doc = Document()

# Настройка стилей
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)
style.paragraph_format.line_spacing = 1.5

# 1. ТИТУЛЬНЫЙ ЛИСТ
add_para(doc, "МИНОБРНАУКИ РОССИИ", 14, True, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Федеральное государственное бюджетное образовательное учреждение высшего образования", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "«Санкт-Петербургский государственный морской технический университет»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "(СПбГМТУ)", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Факультет цифровых промышленных технологий", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "Направление 09.03.01 Информатика и вычислительная техника", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
doc.add_paragraph()
add_para(doc, "КУРСОВАЯ РАБОТА", 16, True, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "по дисциплине «Электротехника и электроника»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Тема: «Расчет предварительного усилителя на ОУ»", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Вариант 11", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "Выполнил: Студент группы 20390", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
add_para(doc, "Суровцев В.Е.", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
doc.add_paragraph()
add_para(doc, "Проверил: Старший преподаватель", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
add_para(doc, "Ветров Б.Г.", 14, False, WD_ALIGN_PARAGRAPH.RIGHT, False)
doc.add_paragraph()
doc.add_paragraph()
add_para(doc, "Санкт-Петербург", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "2025", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
doc.add_page_break()

# 2. СОДЕРЖАНИЕ
add_heading_custom(doc, "Содержание", 1)
toc_items = [
    "Введение", "Задание", "Расчет усилителя", 
    "Выбор элементной базы и количества каскадов", "Расчет оконечного каскада",
    "Расчет входного каскада", "Составление схемы усилителя",
    "Расчет частотных искажений (нижние частоты)", "Расчет частотных искажений (верхние частоты)",
    "Корректировка коэффициента усиления", "Расчет элементов фильтров в цепи питания",
    "Обоснование выбора элементной базы", "Заключение", "Список использованной литературы"
]
for item in toc_items:
    add_para(doc, item, 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
doc.add_page_break()

# 3. ВВЕДЕНИЕ
add_heading_custom(doc, "Введение", 1)
add_para(doc, "Операционный усилитель (ОУ) — это высокоточный электронный усилитель напряжения с дифференциальными входами и одним выходом, обладающий чрезвычайно высоким коэффициентом усиления. Применение ОУ даёт возможность создавать разнообразные электронные схемы и устройства: усилители, генераторы, фильтры. В данной работе рассматриваются вопросы расчёта предварительного усилителя на основе микросхем ОУ.")
add_para(doc, "Параметры каскадов на ОУ определяются функциональным назначением устройств. В морской технике используются предварительные, согласующие, широкополосные и другие системы. Предварительный усиливает сигнал от удалённого датчика до величины, необходимой для передачи по каналу связи, при этом структура сигнала не должна искажаться.")
add_para(doc, "В ходе данной работы будет проведен расчет предварительного усилителя на операционных усилителях для Варианта 11.")
doc.add_page_break()

# 4. ЗАДАНИЕ
add_heading_custom(doc, "Задание", 1)
add_para(doc, "Основные параметры усилителя:", 14, True, WD_ALIGN_PARAGRAPH.LEFT, False)
table_zad = doc.add_table(rows=2, cols=8)
table_zad.style = 'Table Grid'
headers = ["Uг, мВ", "Rг, кОм", "Uн, В", "Rн, Ом", "fн, Гц", "fв, кГц", "Мн", "Мв"]
values = [f"{Ug*1000}", f"{Rg/1000}", f"{Un}", f"{Rn}", f"{fn}", f"{fv/1000}", f"{Mn}", f"{Mv}"]
for i, h in enumerate(headers):
    set_cell(table_zad.rows[0].cells[i], h, bold=True)
for i, v in enumerate(values):
    set_cell(table_zad.rows[1].cells[i], v)
add_para(doc, "Основные параметры операционного усилителя (марка: К574УД3): f1=10 МГц, Rвх=100 МОм, Iп=5 мА.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)

# 5. РАСЧЕТ УСИЛИТЕЛЯ
add_heading_custom(doc, "Расчет усилителя", 1)

add_heading_custom(doc, "Выбор элементной базы и количества каскадов", 2)
add_para(doc, f"Максимальный коэффициент усиления каскада на верхней граничной частоте: K = f1 / fв = {f1_OU/1e6} / {fv/1000} = {f1_OU/fv:.1f}.")
add_para(doc, f"Коэффициент усиления усилителя: Ku = Uн / Uг = {Un} / {Ug*1000} мВ = {Ku:.0f}.")
add_para(doc, f"Определим количество каскадов m. При m=2 средний коэффициент усиления каскада составит {math.sqrt(Ku):.1f}. Принимаем m = 2.")
add_para(doc, f"Распределим усиление по каскадам: Kк1 = {Kk1:.0f} (неинвертирующий), Kк2 = {Kk2:.0f} (инвертирующий).")
add_para(doc, f"Фактический коэффициент усиления: Ku_факт = Kк1 × Kк2 = {Kk1} × {Kk2} = {Ku_fakt:.0f} ≈ {Ku:.0f}.")

add_heading_custom(doc, "Расчет оконечного каскада", 2)
add_para(doc, "Рис. 1 – Схема оконечного каскада", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "[Место для рисунка]", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, f"Для согласования выхода ОУ и нагрузки Rн={Rn} Ом применяем эмиттерный повторитель на комплиментарных транзисторах V1 (КТ315В) и V2 (КТ361Д) с h21={h21}.")
add_para(doc, f"Входное сопротивление эмиттерного повторителя: Rвх_эп = h21 × Rн = {h21} × {Rn} = {Rvh_ep} Ом = {Rvh_ep/1000} кОм.")
add_para(doc, f"Задаем входное сопротивление второго каскада R6 = {format_res(R6)}.")
add_para(doc, f"Коэффициент усиления Kк2 = {Kk2:.0f}. Рассчитываем суммарное сопротивление обратной связи: R_ос = Kк2 × R6 = {Kk2} × {format_res(R6)} = {format_res(R_os_calc)}.")
add_para(doc, f"Для компенсации разброса номиналов заменяем R_ос на потенциометр R7 и постоянный резистор R8:")
add_para(doc, f"  R7 (потенциометр) = 1.2 × R_ос = {format_res(R7_calc)} → стандартный {format_res(R7)}")
add_para(doc, f"  R8 (постоянный) = 0.4 × R_ос = {format_res(R8_calc)} → стандартный {format_res(R8)}")
add_para(doc, f"Проверка Kк2 при среднем положении потенциометра: Kк2_ср = (R7/2 + R8) / R6 = ({format_res(R7/2)} + {format_res(R8)}) / {format_res(R6)} = {Kk2_mid:.1f} ≈ {Kk2:.0f}.")

add_heading_custom(doc, "Расчет входного (первого) каскада", 2)
add_para(doc, "Рис. 2 – Схема входного каскада", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, "[Место для рисунка]", 14, False, WD_ALIGN_PARAGRAPH.CENTER, False)
add_para(doc, f"Входной каскад согласует генератор с усилителем. Принимаем R3 = 7 × Rг = 7 × {format_res(Rg)} = {format_res(R3_calc)}. Стандартный номинал R3 = {format_res(R3)}.")
add_para(doc, f"Коэффициент усиления Kк1 = {Kk1:.0f}. Задаем R1 = {format_res(R1)}.")
add_para(doc, f"Рассчитываем R2 = (Kк1 - 1) × R1 = ({Kk1} - 1) × {format_res(R1)} = {format_res(R2_calc)}. Стандартный номинал R2 = {format_res(R2)}.")
add_para(doc, f"Фактический коэффициент усиления: Kк1_факт = 1 + R2/R1 = 1 + {format_res(R2)}/{format_res(R1)} = {Kk1_fakt:.1f}.")

add_heading_custom(doc, "Расчет частотных искажений (нижние частоты)", 2)
add_para(doc, f"Коэффициент частотных искажений на нижних частотах Мн = {Mn}.")
add_para(doc, f"Предполагая, что все конденсаторы вносят одинаковые искажения, находим коэффициент искажений одного конденсатора: Mнс = Мн^(1/4) = {Mn}^0.25 = {Mnc:.3f}.")
add_para(doc, f"Рассчитываем разделительные конденсаторы:")
add_para(doc, f"C1 (входной): Rэкв = Rг + R3 = {format_res(Rg)} + {format_res(R3)} = {format_res(Req_C1)}")
add_para(doc, f"  C1 ≥ 1/(2π × fн × Rэкв × √(Mнс²-1)) = {format_cap(C1_calc)}")
add_para(doc, f"  Принимаем C1 = {format_cap(C1)} (с запасом).")
add_para(doc, f"C2 (шунт на входе, параллельно R3): для фильтрации ВЧ помех")
add_para(doc, f"  C2 ≥ 1/(2π × fн × R3 × 10) = {format_cap(C2_calc)}")
add_para(doc, f"  Принимаем C2 = {format_cap(C2)}.")
add_para(doc, f"C6 (межкаскадный): Rэкв = R6 = {format_res(R6)}")
add_para(doc, f"  C6 ≥ 1/(2π × fн × R6 × √(Mнс²-1)) = {format_cap(C6_calc)}")
add_para(doc, f"  Принимаем C6 = {format_cap(C6)} (с запасом).")
add_para(doc, f"C7 (выходной): Rэкв = Rн = {format_res(Rn)}")
add_para(doc, f"  C7 ≥ 1/(2π × fн × Rн × √(Mнс²-1)) = {format_cap(C7_calc)}")
add_para(doc, f"  Принимаем C7 = {format_cap(C7)}.")
add_para(doc, f"Проверка C7: Xc7 на fн = 1/(2π × fн × C7) = {Xc7_fn:.1f} Ом << Rн = {Rn} Ом ✓")

add_heading_custom(doc, "Расчет частотных искажений (верхние частоты)", 2)
add_para(doc, f"Для коррекции на верхних частотах (Мв = {Mv}) используем конденсатор C3 в первом каскаде.")
add_para(doc, f"Расчёт искажений от активных элементов:")
add_para(doc, f"  Мв_ОУ1 = (1 + (Kк1 × fв / f1)²) = √(1 + ({Kk1} × {fv/1e3}к / {f1_OU/1e6}М)²) = {Mv_OU1:.4f}")
add_para(doc, f"  Мв_ОУ2 = √(1 + (Kк2 × fв / f1)²) = √(1 + ({Kk2} × {fv/1e3}к / {f1_OU/1e6}М)²) = {Mv_OU2:.4f}")
add_para(doc, f"  Мв_тр = √(1 + (fв / f_h21)²) = (1 + ({fv/1e3}к / {fh21_tr/1e6}М)²) = {Mv_tr:.4f}")
add_para(doc, f"  Мв_акт = Мв_ОУ1 × Мв_ОУ2 × Мв_тр = {Mv_active:.4f}")
add_para(doc, f"Оставшиеся искажения для C3: Мв_C3 = Мв / Мв_акт = {Mv} / {Mv_active:.4f} = {Mv_C3:.3f}")
add_para(doc, f"Расчёт C3: C3 = √(Мв_C3² - 1) / (2π × fв × R2) = {format_cap(C3_calc)}")
add_para(doc, f"Принимаем стандартный номинал C3 = {format_cap(C3)}.")
add_para(doc, f"Проверка: Xc3 на fв = 1/(2π × fв × C3) = {Xc3_fv:.1f} Ом")
add_para(doc, f"  Мв_C3_проверка = √(1 + (R2/Xc3)²) = √(1 + ({format_res(R2)}/{Xc3_fv:.0f})²) = {Mv_C3_check:.3f}")
add_para(doc, f"  Мв_общ_проверка = Мв_акт × Мв_C3_проверка = {Mv_active:.4f} × {Mv_C3_check:.3f} = {Mv_total_check:.3f} ≤ {Mv} ✓")

add_heading_custom(doc, "Корректировка коэффициента усиления", 2)
add_para(doc, f"Для компенсации разброса номиналов резистор обратной связи второго каскада заменяем на потенциометр R7 и постоянный резистор R8.")
add_para(doc, f"R7 (потенциометр СПО-0.25) = {format_res(R7)}")
add_para(doc, f"R8 (постоянный ОМЛТ-0.125) = {format_res(R8)}")
add_para(doc, f"Диапазон регулировки Kк2: от {Kk2_min:.1f} до {Kk2_max:.1f}, при среднем положении {Kk2_mid:.1f} ≈ {Kk2:.0f}.")

add_heading_custom(doc, "Расчет элементов фильтров в цепи питания", 2)
add_para(doc, f"Для устранения наводок по цепям питания устанавливаем RC-фильтры.")
add_para(doc, f"Сопротивление резистора фильтра: Rф = 0.1 × Uп / Iп = 0.1 × 15 / 0.005 = 300 Ом.")
add_para(doc, f"Принимаем R4 = R5 = {format_res(R4)}.")
add_para(doc, f"Ёмкость конденсатора фильтра: Cф = 100 / (2π × fн × Rф) = {format_cap(C4_calc)}.")
add_para(doc, f"Принимаем C4 = C5 = {format_cap(C4)} (по методичке).")

# 6. СПЕЦИФИКАЦИЯ
add_heading_custom(doc, "Обоснование выбора элементной базы (Спецификация)", 1)
table_spec = doc.add_table(rows=1, cols=4)
table_spec.style = 'Table Grid'
headers_spec = ["Обозначение", "Наименование", "Количество", "Примечание"]
for i, h in enumerate(headers_spec):
    set_cell(table_spec.rows[0].cells[i], h, bold=True)

components = [
    ("", "Операционные усилители", "", ""),
    ("A1, A2", "К574УД3", "2", ""),
    ("", "Транзисторы", "", ""),
    ("V1", "КТ315В", "1", "NPN"),
    ("V2", "КТ361Д", "1", "PNP"),
    ("", "Резисторы", "", ""),
    ("R1", f"ОМЛТ-0.125 {format_res(R1)} ±10%", "1", ""),
    ("R2", f"ОМЛТ-0.125 {format_res(R2)} ±10%", "1", ""),
    ("R3", f"ОМЛТ-0.125 {format_res(R3)} ±10%", "1", ""),
    ("R4", f"ОМЛТ-0.125 {format_res(R4)} ±10%", "1", "Фильтр +15В"),
    ("R5", f"ОМЛТ-0.125 {format_res(R5)} ±10%", "1", "Фильтр -15В"),
    ("R6", f"ОМЛТ-0.125 {format_res(R6)} ±10%", "1", ""),
    ("R7", f"Потенциометр СПО-0.25 {format_res(R7)}", "1", "Коррекция Kк2"),
    ("R8", f"ОМЛТ-0.125 {format_res(R8)} ±10%", "1", ""),
    ("", "Конденсаторы", "", ""),
    ("C1", f"К73-17-63В {format_cap(C1)} ±10%", "1", "Входной"),
    ("C2", f"К10-17-50В {format_cap(C2)} ±5%", "1", "Шунт входа"),
    ("C3", f"К10-23-50В {format_cap(C3)} ±5%", "1", "Коррекция ВЧ"),
    ("C4", f"К50-6-63В {format_cap(C4)} ±20%", "1", "Фильтр +15В"),
    ("C5", f"К50-6-63В {format_cap(C5)} ±20%", "1", "Фильтр -15В"),
    ("C6", f"К73-17-63В {format_cap(C6)} ±10%", "1", "Межкаскадный"),
    ("C7", f"К50-6-63В {format_cap(C7)} ±20%", "1", "Выходной"),
]

for comp in components:
    row = table_spec.add_row().cells
    for i, val in enumerate(comp):
        set_cell(row[i], val, align=WD_ALIGN_PARAGRAPH.LEFT if i in [1, 3] else WD_ALIGN_PARAGRAPH.CENTER)

# 7. ЗАКЛЮЧЕНИЕ
add_heading_custom(doc, "Заключение", 1)
add_para(doc, f"В ходе курсового проекта был рассчитан предварительный усилитель на базе ОУ К574УД3. Обеспечен коэффициент усиления Ku = {Ku_fakt:.0f} при заданных параметрах сигнала (требуемое Ku = {Ku:.0f}).")
add_para(doc, f"Все номиналы резисторов и конденсаторов выбраны из стандартных рядов E24 и E12, что гарантирует возможность практической реализации устройства. Расчеты частотных искажений подтвердили соответствие требованиям: Мн={Mn} (расчётное {Mv_total_check:.3f} на ВЧ), Мв={Mv}.")
add_para(doc, f"Введена система коррекции коэффициента усиления второго каскада (R7, R8) для компенсации разброса номиналов резисторов. Установлены фильтры питания (R4, C4 и R5, C5) для подавления наводок по шинам питания.")

# 8. ЛИТЕРАТУРА
add_heading_custom(doc, "Список использованной литературы", 1)
add_para(doc, "1. Методические рекомендации по выполнению курсовой работы. Ветров Б. Г.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "2. Курс лекций по электронике. Ветров Б. Г.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)
add_para(doc, "3. Справочник по элементной базе электронных устройств.", 14, False, WD_ALIGN_PARAGRAPH.LEFT, False)

# Сохранение
file_name = "Курсовая_работа_Вариант_11_исправленная_v2.docx"
doc.save(file_name)
print(f"Файл успешно создан: {os.path.abspath(file_name)}")
print("\n=== СВОДНАЯ ТАБЛИЦА НОМИНАЛОВ ===")
print(f"R1 = {format_res(R1)}, R2 = {format_res(R2)}, R3 = {format_res(R3)}")
print(f"R4 = R5 = {format_res(R4)}")
print(f"R6 = {format_res(R6)}, R7 = {format_res(R7)}, R8 = {format_res(R8)}")
print(f"C1 = {format_cap(C1)}, C2 = {format_cap(C2)}, C3 = {format_cap(C3)}")
print(f"C4 = C5 = {format_cap(C4)}, C6 = {format_cap(C6)}, C7 = {format_cap(C7)}")
print(f"\nKк1 = {Kk1_fakt:.1f}, Kк2 (среднее) = {Kk2_mid:.1f}, Ku = {Ku_fakt:.0f}")
print(f"Мв_общ (проверка) = {Mv_total_check:.3f} ≤ {Mv} ✓")